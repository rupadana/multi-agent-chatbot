from datetime import datetime, timezone
import io

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlmodel import Session, select

from ..database import get_session
from ..deps import get_current_user
from ..models import Agent, Document, User
from ..schemas import DocumentCreate, DocumentRead
from .agents import get_agent_or_404

router = APIRouter(prefix="/api/agents/{agent_id}/knowledge", tags=["knowledge"])


@router.get("", response_model=list[DocumentRead])
def list_documents(
    agent_id: int,
    session: Session = Depends(get_session),
    user: User = Depends(get_current_user),
):
    get_agent_or_404(agent_id, session, user)
    docs = session.exec(
        select(Document)
        .where(Document.agent_id == agent_id)
        .order_by(Document.created_at.desc())
    ).all()
    return docs


@router.post("", response_model=DocumentRead, status_code=201)
def add_document(
    agent_id: int,
    payload: DocumentCreate,
    session: Session = Depends(get_session),
    user: User = Depends(get_current_user),
):
    agent: Agent = get_agent_or_404(agent_id, session, user)
    doc = Document(agent_id=agent.id, title=payload.title, content=payload.content)
    session.add(doc)
    # Sentuh updated_at agent agar urutan tetap relevan.
    agent.updated_at = datetime.now(timezone.utc)
    session.add(agent)
    session.commit()
    session.refresh(doc)
    return doc


@router.delete("/{document_id}", status_code=204)
def delete_document(
    agent_id: int,
    document_id: int,
    session: Session = Depends(get_session),
    user: User = Depends(get_current_user),
):
    get_agent_or_404(agent_id, session, user)
    doc = session.get(Document, document_id)
    if doc is None or doc.agent_id != agent_id:
        raise HTTPException(status_code=404, detail="Dokumen tidak ditemukan")
    session.delete(doc)
    session.commit()


@router.put("/{document_id}", response_model=DocumentRead)
def update_document(
    agent_id: int,
    document_id: int,
    payload: DocumentCreate,
    session: Session = Depends(get_session),
    user: User = Depends(get_current_user),
):
    agent = get_agent_or_404(agent_id, session, user)
    doc = session.get(Document, document_id)
    if doc is None or doc.agent_id != agent_id:
        raise HTTPException(status_code=404, detail="Dokumen tidak ditemukan")

    doc.title = payload.title
    doc.content = payload.content
    session.add(doc)

    # Touch agent update timestamp
    agent.updated_at = datetime.now(timezone.utc)
    session.add(agent)

    session.commit()
    session.refresh(doc)
    return doc


def parse_uploaded_file(filename: str, file_content: bytes) -> str:
    ext = filename.split(".")[-1].lower() if "." in filename else ""

    if ext in ("txt", "md", "markdown"):
        try:
            return file_content.decode("utf-8")
        except UnicodeDecodeError:
            return file_content.decode("latin-1")

    elif ext == "docx":
        try:
            import docx
        except ImportError:
            raise HTTPException(
                status_code=500,
                detail="Library python-docx tidak terinstal di server."
            )
        doc_file = io.BytesIO(file_content)
        doc = docx.Document(doc_file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                full_text.append(" | ".join(row_text))
        return "\n".join(full_text)

    elif ext == "xlsx":
        try:
            import openpyxl
        except ImportError:
            raise HTTPException(
                status_code=500,
                detail="Library openpyxl tidak terinstal di server."
            )
        xls_file = io.BytesIO(file_content)
        wb = openpyxl.load_workbook(xls_file, data_only=True)
        sheets_text = []
        for name in wb.sheetnames:
            ws = wb[name]
            sheets_text.append(f"### Sheet: {name}")
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue

            non_empty_rows = [r for r in rows if any(cell is not None for cell in r)]
            if not non_empty_rows:
                continue

            header = non_empty_rows[0]
            header_str = " | ".join(str(val) if val is not None else "" for val in header)
            sep_str = " | ".join("---" for _ in range(len(header)))
            sheets_text.append(f"| {header_str} |")
            sheets_text.append(f"| {sep_str} |")

            for row in non_empty_rows[1:]:
                row_str = " | ".join(str(val) if val is not None else "" for val in row)
                sheets_text.append(f"| {row_str} |")

            sheets_text.append("")
        return "\n".join(sheets_text)

    elif ext == "pdf":
        try:
            import pypdf
        except ImportError:
            raise HTTPException(
                status_code=500,
                detail="Library pypdf tidak terinstal di server."
            )
        pdf_file = io.BytesIO(file_content)
        reader = pypdf.PdfReader(pdf_file)
        pages_text = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages_text.append(text)
        return "\n\n".join(pages_text)

    elif ext == "doc":
        raise HTTPException(
            status_code=400,
            detail="Format .doc (lama) tidak didukung secara langsung. Harap simpan dokumen sebagai .docx dan unggah kembali."
        )

    else:
        raise HTTPException(
            status_code=400,
            detail=f"Format file .{ext} tidak didukung. Harap unggah format txt, md, docx, xlsx, atau pdf."
        )


@router.post("/upload", response_model=DocumentRead, status_code=201)
def upload_document(
    agent_id: int,
    file: UploadFile = File(...),
    session: Session = Depends(get_session),
    user: User = Depends(get_current_user),
):
    agent = get_agent_or_404(agent_id, session, user)

    file_content = file.file.read()
    content_text = parse_uploaded_file(file.filename, file_content)

    original_title = file.filename
    if "." in original_title:
        title = ".".join(original_title.split(".")[:-1])
    else:
        title = original_title

    doc = Document(agent_id=agent.id, title=title, content=content_text)
    session.add(doc)

    # Touch agent update timestamp
    agent.updated_at = datetime.now(timezone.utc)
    session.add(agent)

    session.commit()
    session.refresh(doc)
    return doc


